"""
답변 내보내기 — Excel / Word / PDF
마크다운 텍스트를 받아 각 형식으로 변환해서 다운로드
"""
from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import io, re

router = APIRouter()

class ExportRequest(BaseModel):
    content: str        # 마크다운 텍스트
    title: str = "챗봇 답변"
    format: str = "excel"  # excel | word | pdf

# ── 마크다운 파싱 헬퍼 ─────────────────────────────────────────────────────────
def parse_md_tables(text: str) -> list[dict]:
    """마크다운 텍스트에서 테이블 추출 → [{"title": str, "rows": [[...]]}]"""
    tables = []
    lines  = text.split('\n')
    i, n   = 0, len(lines)

    while i < n:
        # 테이블 앞에 제목이 있으면 추출
        title = ""
        if i > 0:
            prev = lines[i-1].strip()
            if prev and not prev.startswith('|') and not prev.startswith('-'):
                title = re.sub(r'^#+\s*', '', prev).strip('* ')

        # 헤더 행 감지
        if lines[i].strip().startswith('|') and i+1 < n and re.match(r'\s*\|[-:| ]+\|', lines[i+1]):
            header = [c.strip() for c in lines[i].strip('|').split('|')]
            rows   = [header]
            i += 2  # 구분선 건너뜀
            while i < n and lines[i].strip().startswith('|'):
                row = [c.strip() for c in lines[i].strip('|').split('|')]
                rows.append(row)
                i += 1
            tables.append({"title": title or f"표 {len(tables)+1}", "rows": rows})
        else:
            i += 1
    return tables

def extract_code_blocks(text: str) -> list[dict]:
    """코드 블록 추출 → [{"lang": str, "code": str}]"""
    pattern = r'```(\w*)\n([\s\S]*?)```'
    return [{"lang": m.group(1), "code": m.group(2).strip()}
            for m in re.finditer(pattern, text)]

def clean_md(text: str) -> str:
    """마크다운 기호 제거 → 순수 텍스트"""
    text = re.sub(r'#{1,6}\s*', '', text)
    text = re.sub(r'\*{1,3}([^*]+)\*{1,3}', r'\1', text)
    text = re.sub(r'`{3}[^\n]*\n([\s\S]*?)`{3}', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    text = re.sub(r'^\s*[-*+]\s+', '• ', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*\|.*\|.*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*[-|: ]+$', '', text, flags=re.MULTILINE)
    return re.sub(r'\n{3,}', '\n\n', text).strip()

# ── Excel 생성 ────────────────────────────────────────────────────────────────
def to_excel(content: str, title: str) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    wb.remove(wb.active)

    tables  = parse_md_tables(content)
    codes   = extract_code_blocks(content)
    NAVY    = "0F1B2D"
    BLUE_LT = "DBEAFE"
    BORDER  = Border(
        left=Side(style='thin', color='E2E8F0'),
        right=Side(style='thin', color='E2E8F0'),
        top=Side(style='thin', color='E2E8F0'),
        bottom=Side(style='thin', color='E2E8F0'),
    )

    def add_table_sheet(ws_title, rows):
        ws = wb.create_sheet(title=ws_title[:31])
        if not rows: return
        # 헤더
        for ci, cell_val in enumerate(rows[0], 1):
            cell = ws.cell(row=1, column=ci, value=cell_val)
            cell.font = Font(bold=True, color="FFFFFF", name="맑은 고딕", size=10)
            cell.fill = PatternFill(fill_type="solid", fgColor=NAVY)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = BORDER
        # 데이터
        for ri, row in enumerate(rows[1:], 2):
            fill = PatternFill(fill_type="solid", fgColor=BLUE_LT if ri%2==0 else "FFFFFF")
            for ci, val in enumerate(row, 1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.alignment = Alignment(vertical="center", wrap_text=True)
                cell.font = Font(name="맑은 고딕", size=10)
                cell.fill = fill
                cell.border = BORDER
        # 열 너비 자동
        for col in ws.columns:
            max_len = max((len(str(c.value or "")) for c in col), default=10)
            ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 4, 40)
        ws.row_dimensions[1].height = 22

    # 테이블 시트
    if tables:
        for t in tables:
            add_table_sheet(t["title"], t["rows"])

    # SQL 코드 시트
    sql_codes = [c for c in codes if c["lang"].lower() in ("sql", "")]
    if sql_codes:
        ws = wb.create_sheet(title="SQL")
        ws.cell(row=1, column=1, value="SQL 코드").font = Font(bold=True, size=11, color="FFFFFF", name="맑은 고딕")
        ws.cell(row=1, column=1).fill = PatternFill(fill_type="solid", fgColor=NAVY)
        r = 2
        for c in sql_codes:
            for line in c["code"].split('\n'):
                ws.cell(row=r, column=1, value=line).font = Font(name="Consolas", size=10)
                r += 1
            r += 1
        ws.column_dimensions['A'].width = 70

    # 전체 텍스트 시트
    ws = wb.create_sheet(title="전체 내용", index=0)
    plain = clean_md(content)
    ws.cell(row=1, column=1, value=title).font = Font(bold=True, size=13, name="맑은 고딕", color=NAVY)
    r = 2
    for line in plain.split('\n'):
        ws.cell(row=r, column=1, value=line).font = Font(name="맑은 고딕", size=10)
        r += 1
    ws.column_dimensions['A'].width = 80

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()

# ── Word 생성 ─────────────────────────────────────────────────────────────────
def to_word(content: str, title: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # 제목
    h = doc.add_heading(title, level=1)
    h.runs[0].font.color.rgb = RGBColor(0x0F, 0x1B, 0x2D)

    tables = parse_md_tables(content)
    codes  = extract_code_blocks(content)

    # 섹션별 처리
    for line in content.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        m = re.match(r'^(#{1,4})\s+(.+)', stripped)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2), level=level)
        elif stripped.startswith('```'):
            pass
        elif re.match(r'^\|.+\|', stripped):
            pass
        elif stripped.startswith(('- ', '* ', '• ')):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(stripped.lstrip('-* •').strip()).font.size = Pt(10)
        else:
            # 볼드 처리
            p = doc.add_paragraph()
            parts = re.split(r'\*{1,3}([^*]+)\*{1,3}', stripped)
            bold = False
            for pi, part in enumerate(parts):
                run = p.add_run(part)
                run.font.size = Pt(10)
                run.bold = bool(pi % 2)

    # 테이블 추가
    for t in tables:
        if not t["rows"]: continue
        doc.add_heading(t["title"], level=3)
        tbl = doc.add_table(rows=len(t["rows"]), cols=len(t["rows"][0]))
        tbl.style = 'Table Grid'
        for ri, row in enumerate(t["rows"]):
            for ci, val in enumerate(row):
                cell = tbl.cell(ri, ci)
                cell.text = val
                run = cell.paragraphs[0].runs
                if run and ri == 0:
                    run[0].bold = True
                    run[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # SQL 코드
    sql_codes = [c for c in codes if c["lang"].lower() in ("sql", "")]
    for c in sql_codes:
        doc.add_heading("SQL", level=3)
        p = doc.add_paragraph(c["code"])
        p.runs[0].font.name = "Consolas"
        p.runs[0].font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ── PDF 생성 ──────────────────────────────────────────────────────────────────
def to_pdf(content: str, title: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Preformatted
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os

    # 한글 폰트 (시스템 폰트 탐색)
    font_paths = [
        "/System/Library/Fonts/AppleSDGothicNeo.ttc",
        "/Library/Fonts/NanumGothic.ttf",
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
    ]
    font_name = "Helvetica"
    for fp in font_paths:
        if os.path.exists(fp):
            try:
                pdfmetrics.registerFont(TTFont("Korean", fp))
                font_name = "Korean"
                break
            except Exception:
                pass

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=20*mm, rightMargin=20*mm,
                            topMargin=20*mm, bottomMargin=20*mm)
    styles = getSampleStyleSheet()
    NAVY = colors.HexColor('#0F1B2D')

    title_style = ParagraphStyle('Title', fontName=font_name, fontSize=16,
                                  textColor=NAVY, spaceAfter=12)
    h2_style    = ParagraphStyle('H2', fontName=font_name, fontSize=13,
                                  textColor=NAVY, spaceBefore=10, spaceAfter=6, leading=16)
    body_style  = ParagraphStyle('Body', fontName=font_name, fontSize=10,
                                  leading=16, spaceAfter=4)
    code_style  = ParagraphStyle('Code', fontName='Courier', fontSize=8,
                                  backColor=colors.HexColor('#F8F8F8'), leading=12)

    story = [Paragraph(title, title_style), Spacer(1, 8)]

    tables = parse_md_tables(content)
    codes  = extract_code_blocks(content)
    plain  = clean_md(content)

    for line in plain.split('\n'):
        if not line.strip(): continue
        if re.match(r'^#{1,4}\s', line):
            text = re.sub(r'^#+\s*', '', line)
            story.append(Paragraph(text, h2_style))
        elif line.startswith('•'):
            story.append(Paragraph(f"&bull; {line[1:].strip()}", body_style))
        else:
            story.append(Paragraph(line, body_style))

    for t in tables:
        if not t["rows"]: continue
        story.append(Spacer(1, 6))
        story.append(Paragraph(t["title"], h2_style))
        tbl_data = t["rows"]
        tbl = Table(tbl_data, repeatRows=1)
        tbl.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), NAVY),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,-1), font_name),
            ('FONTSIZE', (0,0), (-1,-1), 9),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#EFF6FF')]),
            ('PADDING', (0,0), (-1,-1), 5),
        ]))
        story.append(tbl)

    for c in [c for c in codes if c["lang"].lower() in ("sql","")]:
        story.append(Spacer(1, 6))
        story.append(Paragraph("SQL", h2_style))
        story.append(Preformatted(c["code"], code_style))

    doc.build(story)
    return buf.getvalue()

# ── 엔드포인트 ────────────────────────────────────────────────────────────────
@router.post("/export")
async def export_answer(req: ExportRequest):
    fmt = req.format.lower()

    if fmt == "excel":
        data = to_excel(req.content, req.title)
        return StreamingResponse(
            io.BytesIO(data),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="answer.xlsx"'}
        )
    elif fmt == "word":
        data = to_word(req.content, req.title)
        return StreamingResponse(
            io.BytesIO(data),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="answer.docx"'}
        )
    elif fmt == "pdf":
        data = to_pdf(req.content, req.title)
        return StreamingResponse(
            io.BytesIO(data),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="answer.pdf"'}
        )
    else:
        return {"error": "지원하지 않는 형식"}
