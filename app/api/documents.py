"""
사용자 문서 업로드 → 텍스트 추출 → 청크 분할 → 임베딩 → RAG 포함

문서 타입별 메타데이터 처리:
- PDF   : 페이지 번호(page_no)
- Excel : 시트명(sheet_name), 행 범위(row_range)
- Word  : 섹션/제목(section)
- TXT/MD: 섹션(section, 제목 기준)
- HWP   : 텍스트 추출 (구조 메타 없음)
"""
from fastapi import APIRouter, UploadFile, File, HTTPException, BackgroundTasks
from typing import List
import io, os, re, json, requests
from ..core.database import db_cursor
from ..core.config import settings

router = APIRouter()

# ── 임베딩 ───────────────────────────────────────────────────────────────────
def _embed(text: str):
    try:
        r = requests.post(settings.ollama_url,
                          json={"model": settings.embed_model, "prompt": text[:2000]},
                          timeout=30)
        return r.json()["embedding"]
    except Exception:
        return None

# ══════════════════════════════════════════════════════════════════════════════
# 타입별 추출 — (content, metadata) 튜플 리스트 반환
# metadata = {"page_no", "sheet_name", "row_range", "section"}
# ══════════════════════════════════════════════════════════════════════════════

def _extract_pdf(data: bytes) -> list[dict]:
    """PDF → 페이지 단위 청크 (큰 페이지는 2분할)"""
    import pdfplumber
    chunks = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page_no, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            text = text.strip()
            if not text:
                continue
            # 한 페이지가 너무 길면 단락 기준으로 분할
            if len(text) > 1200:
                parts = _split_by_paragraph(text, max_size=1000)
                for part in parts:
                    if part.strip():
                        chunks.append({"content": part, "page_no": page_no,
                                       "sheet_name": None, "row_range": None, "section": None})
            else:
                chunks.append({"content": text, "page_no": page_no,
                               "sheet_name": None, "row_range": None, "section": None})
    return chunks


def _extract_docx(data: bytes) -> list[dict]:
    """Word → 섹션(제목) 단위 청크"""
    from docx import Document
    doc = Document(io.BytesIO(data))
    chunks, current_section, current_buf = [], "본문", []

    def flush():
        if current_buf:
            text = "\n".join(current_buf).strip()
            if len(text) > 30:
                parts = _split_by_paragraph(text, max_size=800)
                for p in parts:
                    if p.strip():
                        chunks.append({"content": p, "section": current_section,
                                       "page_no": None, "sheet_name": None, "row_range": None})
            current_buf.clear()

    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        style = para.style.name or ""
        if any(style.startswith(h) for h in ("Heading", "제목")):
            flush()
            current_section = para.text.strip()
        else:
            current_buf.append(para.text.strip())
    flush()

    # 표 처리
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        header = " | ".join(rows[0])
        body   = "\n".join(" | ".join(r) for r in rows[1:])
        text   = f"[표]\n{header}\n{body}"
        chunks.append({"content": text, "section": "표",
                       "page_no": None, "sheet_name": None, "row_range": None})
    return chunks


def _extract_xlsx(data: bytes) -> list[dict]:
    """Excel → 시트·섹션 단위 청크 (빈 행으로 섹션 구분)"""
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    chunks = []

    for ws in wb.worksheets:
        rows_data = list(ws.iter_rows(values_only=True))
        if not rows_data:
            continue

        # 헤더 행 탐지 (첫 행)
        header = [str(v).strip() if v is not None else "" for v in rows_data[0]]

        # 빈 행 기준으로 섹션 분할
        sections, current, start_row = [], [], 1
        for ri, row in enumerate(rows_data[1:], 2):
            if all(v is None for v in row):
                if current:
                    sections.append((start_row, ri - 1, current))
                    current = []
                    start_row = ri + 1
            else:
                current.append(row)
        if current:
            sections.append((start_row, len(rows_data), current))

        # 섹션이 없으면 전체를 하나로
        if not sections:
            sections = [(1, len(rows_data), rows_data[1:])]

        for s_start, s_end, rows in sections:
            lines = [" | ".join(header)]
            for row in rows[:50]:  # 한 청크 최대 50행
                line = " | ".join(str(v).strip() if v is not None else "-" for v in row)
                if line.replace("|", "").strip():
                    lines.append(line)
            if len(lines) > 1:
                chunks.append({
                    "content": "\n".join(lines),
                    "sheet_name": ws.title,
                    "row_range": f"{s_start}-{s_end}",
                    "page_no": None,
                    "section": None,
                })
    return chunks


def _extract_txt(data: bytes, filename: str) -> list[dict]:
    """TXT/MD → 제목(#) 기준 섹션 청크"""
    text = data.decode("utf-8", errors="replace")
    ext = os.path.splitext(filename)[1].lower()
    chunks, current_section, current_buf = [], "본문", []

    def flush():
        if current_buf:
            content = "\n".join(current_buf).strip()
            parts = _split_by_paragraph(content, max_size=800)
            for p in parts:
                if p.strip():
                    chunks.append({"content": p, "section": current_section,
                                   "page_no": None, "sheet_name": None, "row_range": None})
            current_buf.clear()

    for line in text.split("\n"):
        m = re.match(r'^(#{1,4})\s+(.+)', line)
        if m:
            flush()
            current_section = m.group(2).strip()
        else:
            if line.strip():
                current_buf.append(line)
    flush()
    return chunks


def _extract_hwp(data: bytes) -> list[dict]:
    """HWP — 텍스트만 추출 (구조 정보 없음)"""
    try:
        # HWP5 바이너리에서 한글 텍스트 패턴 추출
        text = data.decode("utf-8", errors="replace")
        text = re.sub(r'[^가-힣ㄱ-ㅎㅏ-ㅣ\w\s.,;:!?()\-/\n]', ' ', text)
        text = re.sub(r'\s{3,}', '\n\n', text).strip()
    except Exception:
        text = ""

    if not text:
        return []

    parts = _split_by_paragraph(text, max_size=800)
    return [{"content": p, "page_no": None, "sheet_name": None,
             "row_range": None, "section": None}
            for p in parts if len(p.strip()) > 30]


# ── 단락 기준 분할 유틸 ───────────────────────────────────────────────────────
def _split_by_paragraph(text: str, max_size: int = 800, overlap_chars: int = 150) -> list[str]:
    """
    단락(\n\n) 기준 분할 + 문자 단위 오버랩
    - max_size 초과 시 새 청크 시작
    - 이전 청크 끝 overlap_chars만큼을 다음 청크 앞에 붙여 문맥 연속성 확보
    """
    paragraphs = [p.strip() for p in re.split(r'\n{2,}', text) if p.strip()]
    if not paragraphs:
        return []

    chunks, buf, buf_len = [], [], 0

    for para in paragraphs:
        plen = len(para)
        if buf_len + plen > max_size and buf:
            chunk_text = "\n\n".join(buf)
            chunks.append(chunk_text)
            # 오버랩: 이전 청크의 끝부분을 다음 청크 시작에 포함
            overlap_text = chunk_text[-overlap_chars:] if len(chunk_text) > overlap_chars else chunk_text
            buf = [overlap_text] if overlap_text.strip() else []
            buf_len = len(overlap_text) if buf else 0
        buf.append(para)
        buf_len += plen

    if buf:
        chunks.append("\n\n".join(buf))

    # 너무 긴 단락(단락 자체가 max_size 초과)은 문장 단위로 추가 분할
    final = []
    for chunk in chunks:
        if len(chunk) <= max_size * 1.5:
            final.append(chunk)
        else:
            sentences = re.split(r'(?<=[.!?。])\s+', chunk)
            sub, sub_len = [], 0
            for sent in sentences:
                slen = len(sent)
                if sub_len + slen > max_size and sub:
                    final.append(" ".join(sub))
                    sub, sub_len = [], 0
                sub.append(sent)
                sub_len += slen
            if sub:
                final.append(" ".join(sub))
    return [c for c in final if len(c.strip()) > 30]


# ── 파일 종류 판별 후 추출 ────────────────────────────────────────────────────
def _extract_csv(data: bytes) -> list[dict]:
    """CSV → Excel처럼 표 구조로 처리"""
    import csv
    text = data.decode("utf-8", errors="replace")
    reader = list(csv.reader(text.splitlines()))
    if not reader:
        return []

    header = [str(v).strip() for v in reader[0]]
    chunks, buf_rows, start_row = [], [], 1

    for ri, row in enumerate(reader[1:], 2):
        buf_rows.append(row)
        if len(buf_rows) >= 50:  # 50행마다 청크
            lines = [" | ".join(header)]
            for r in buf_rows:
                line = " | ".join(str(v).strip() for v in r)
                if line.replace("|", "").strip():
                    lines.append(line)
            if len(lines) > 1:
                chunks.append({
                    "content": "\n".join(lines),
                    "sheet_name": "CSV",
                    "row_range": f"{start_row}-{ri}",
                    "page_no": None, "section": None,
                })
            start_row = ri + 1
            buf_rows = []

    if buf_rows:
        lines = [" | ".join(header)]
        for r in buf_rows:
            line = " | ".join(str(v).strip() for v in r)
            if line.replace("|", "").strip():
                lines.append(line)
        if len(lines) > 1:
            chunks.append({
                "content": "\n".join(lines),
                "sheet_name": "CSV",
                "row_range": f"{start_row}-{start_row+len(buf_rows)-1}",
                "page_no": None, "section": None,
            })
    return chunks


def extract_chunks(data: bytes, filename: str) -> list[dict]:
    ext = os.path.splitext(filename)[1].lower()
    if   ext == ".pdf":               return _extract_pdf(data)
    elif ext == ".docx":              return _extract_docx(data)  # .doc(구버전)은 지원 불가
    elif ext in (".xlsx", ".xls"):    return _extract_xlsx(data)
    elif ext == ".csv":               return _extract_csv(data)
    elif ext in (".txt", ".md"):      return _extract_txt(data, filename)
    elif ext == ".hwp":               return _extract_hwp(data)
    else:
        raise HTTPException(400, f"지원하지 않는 형식: {ext}")


# ── 백그라운드 처리 파이프라인 ─────────────────────────────────────────────────
def process_document(doc_id: str, data: bytes, filename: str):
    """추출 → 임베딩 → DB 저장 (백그라운드)"""
    try:
        chunks = extract_chunks(data, filename)
        if not chunks:
            with db_cursor() as cur:
                cur.execute(
                    "UPDATE user_documents SET status='error', error_msg=%s, updated_at=NOW() WHERE id=%s",
                    ("텍스트를 추출할 수 없습니다.", doc_id)
                )
            return

        with db_cursor() as cur:
            for i, chunk in enumerate(chunks):
                vec = _embed(chunk["content"])
                vec_str = ("[" + ",".join(map(str, vec)) + "]") if vec else None
                cur.execute("""
                    INSERT INTO user_doc_chunks
                        (doc_id, chunk_no, page_no, sheet_name, row_range, section, content, embedding)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s::vector)
                """, (
                    doc_id, i,
                    chunk.get("page_no"),
                    chunk.get("sheet_name"),
                    chunk.get("row_range"),
                    chunk.get("section"),
                    chunk["content"],
                    vec_str,
                ))
            cur.execute(
                "UPDATE user_documents SET status='ready', total_chunks=%s, updated_at=NOW() WHERE id=%s",
                (len(chunks), doc_id)
            )

    except Exception as e:
        with db_cursor() as cur:
            cur.execute(
                "UPDATE user_documents SET status='error', error_msg=%s, updated_at=NOW() WHERE id=%s",
                (str(e)[:200], doc_id)
            )


# ── API 엔드포인트 ─────────────────────────────────────────────────────────────
ALLOWED_EXT = {".pdf", ".docx", ".xlsx", ".xls", ".txt", ".md", ".csv", ".hwp"}
# 지원 불가: .doc (구버전 Word 바이너리 — python-docx 미지원)

@router.post("/documents/upload")
async def upload_document(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in ALLOWED_EXT:
        raise HTTPException(400, f"지원 형식: {', '.join(sorted(ALLOWED_EXT))}")

    data = await file.read()
    if len(data) == 0:
        raise HTTPException(400, "빈 파일입니다.")
    if len(data) > 50 * 1024 * 1024:
        raise HTTPException(400, "파일 크기는 50MB 이하여야 합니다.")

    with db_cursor() as cur:
        cur.execute(
            "INSERT INTO user_documents (filename, file_type, file_size) VALUES (%s, %s, %s) RETURNING id",
            (file.filename, ext, len(data))
        )
        doc_id = str(cur.fetchone()["id"])

    background_tasks.add_task(process_document, doc_id, data, file.filename)
    return {"id": doc_id, "filename": file.filename, "status": "processing"}


@router.get("/documents")
def list_documents():
    with db_cursor() as cur:
        cur.execute("""
            SELECT d.id, d.filename, d.file_type, d.file_size,
                   d.total_chunks, d.status, d.error_msg, d.created_at,
                   COUNT(c.id) FILTER (WHERE c.embedding IS NOT NULL) AS embedded_chunks
            FROM user_documents d
            LEFT JOIN user_doc_chunks c ON c.doc_id = d.id
            GROUP BY d.id
            ORDER BY d.created_at DESC
        """)
        return [dict(r) for r in cur.fetchall()]


@router.get("/documents/{doc_id}/status")
def get_status(doc_id: str):
    with db_cursor() as cur:
        cur.execute(
            "SELECT id, filename, status, total_chunks, error_msg FROM user_documents WHERE id=%s",
            (doc_id,)
        )
        row = cur.fetchone()
        if not row: raise HTTPException(404, "Not found")
        return dict(row)


@router.delete("/documents/{doc_id}")
def delete_document(doc_id: str):
    with db_cursor() as cur:
        cur.execute("DELETE FROM user_documents WHERE id=%s RETURNING id", (doc_id,))
        if not cur.fetchone(): raise HTTPException(404, "Not found")
    return {"deleted": doc_id}
